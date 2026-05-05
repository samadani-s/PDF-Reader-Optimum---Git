
from flask import Flask, render_template, request, send_from_directory
from flask_socketio import SocketIO
from pdf2docx import Converter
from docx import Document
from pdf2image import convert_from_path
import pytesseract
import os
import time

app = Flask(__name__)
socketio = SocketIO(app)

UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

def log(msg, level='log'):
    socketio.emit('log', {'message': msg, 'level': level})
    time.sleep(1)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/', methods=['POST'])
def upload():
    file = request.files['file']
    filename = file.filename
    pdf_path = os.path.join(UPLOAD_FOLDER, filename)
    docx_path = pdf_path.replace('.pdf', '.docx')
    file.save(pdf_path)
    socketio.start_background_task(convert_pdf, pdf_path, docx_path)
    return '', 204

@app.route('/download/<filename>')
def download_file(filename):
    return send_from_directory(UPLOAD_FOLDER, filename, as_attachment=True)

def docx_has_real_text(docx_path):
    try:
        doc = Document(docx_path)
        for para in doc.paragraphs:
            if para.text.strip():
                return True
        return False
    except:
        return False

def convert_pdf(pdf_path, docx_path):
    log("شروع پردازش فایل...", "info")
    temp_docx = docx_path.replace('.docx', '_temp.docx')

    try:
        log("تلاش برای تبدیل متنی با pdf2docx...", "info")
        cv = Converter(pdf_path)
        cv.convert(temp_docx)
        cv.close()

        if docx_has_real_text(temp_docx):
            os.rename(temp_docx, docx_path)
            log("تبدیل موفق با متن واقعی (pdf2docx)", "success")
        else:
            log("فایل فقط تصویر دارد. اجرای OCR...", "warn")
            os.remove(temp_docx)
            images = convert_from_path(pdf_path)
            doc = Document()
            for i, img in enumerate(images):
                log(f"OCR صفحه {i+1} در حال انجام...", "info")
                text = pytesseract.image_to_string(img, lang='fas')
                doc.add_paragraph(text)
            doc.save(docx_path)
            log("متن با OCR استخراج شد.", "success")

        log("تبدیل کامل شد. فایل Word آماده است.", "success")
        socketio.emit('done', {'path': f'/download/{os.path.basename(docx_path)}'})

    except Exception as e:
        log(f"خطا در تبدیل: {str(e)}", "error")
